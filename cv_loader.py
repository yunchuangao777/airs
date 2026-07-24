from pathlib import Path
import fitz  # PyMuPDF
from docx import Document
from zipfile import ZipFile
from lxml import etree

def extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF."""
    text = []

    try:
        doc = fitz.open(file_path)

        for page in doc:
            text.append(page.get_text())

        doc.close()

    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")

    return "\n".join(text)


def extract_docx_text(file_path: str) -> str:
    texts = []

    # 1. Normal paragraphs and tables
    doc = Document(file_path)

    for p in doc.paragraphs:
        if p.text.strip():
            texts.append(p.text.strip())

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    texts.append(cell.text.strip())

    # 2. Text boxes / shapes / headers / footers from raw XML
    try:
        with ZipFile(file_path) as z:
            xml_files = [
                name for name in z.namelist()
                if name.startswith("word/") and name.endswith(".xml")
            ]

            for xml_file in xml_files:
                xml_content = z.read(xml_file)
                root = etree.fromstring(xml_content)

                namespaces = {
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                }

                for node in root.xpath(".//w:t", namespaces=namespaces):
                    if node.text and node.text.strip():
                        texts.append(node.text.strip())

    except Exception as e:
        print(f"Warning: XML extraction failed for {file_path}: {e}")

    # Remove duplicates while preserving order
    seen = set()
    clean_texts = []

    for t in texts:
        if t not in seen:
            clean_texts.append(t)
            seen.add(t)

    return "\n".join(clean_texts)


def extract_txt_text(file_path: str) -> str:
    """Extract text from TXT."""

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    except UnicodeDecodeError:
        with open(file_path, "r", encoding="gbk") as f:
            return f.read()

    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        return ""


def load_single_cv(file_path: Path) -> dict:
    """
    Return:
    {
        "filename": "...",
        "filepath": "...",
        "text": "..."
    }
    """

    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        text = extract_pdf_text(str(file_path))

    elif suffix == ".docx":
        text = extract_docx_text(str(file_path))

    elif suffix == ".txt":
        text = extract_txt_text(str(file_path))

    else:
        text = ""

    return {
        "filename": file_path.name,
        "filepath": str(file_path),
        "text": text
    }


def load_cv_folder(folder_path: str) -> list:
    """
    Recursively load all CV files under a folder.
    """

    folder = Path(folder_path)

    supported = {".pdf", ".docx", ".txt"}

    results = []

    for file in folder.rglob("*"):

        if file.is_file() and file.suffix.lower() in supported:

            print(f"Loading: {file.name}")

            cv = load_single_cv(file)

            results.append(cv)

    return results